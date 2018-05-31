import argparse

import string
import re

import docx


REGEX = '\([a-zA-Z][^)\[\]=]*[0-9]\)'


EXCLUDE = set(string.punctuation) - set('-')


def find_reference_for_cite_regex(cite_regex):
    for reference in references:
        result = re.findall(cite_regex, reference)
        if result:
            assert len(result) == 1
            return reference


def get_references_as_lines(docx_path):
    doc = docx.Document(docx_path)

    go = False
    result = list()

    for p in doc.paragraphs:
        if p.text == 'References':
            go = True
        elif go and p.text.startswith('Table'):
            go = False

        if go:
            result.append(p.text)
    return result


def get_text_as_string(docx_path):
    doc = docx.Document(docx_path)

    go = True
    result = list()

    for p in doc.paragraphs:
        if p.text == 'References':
            go = False

        if go:
            result.append(p.text)

    return " ".join(result)


class CitationResult(object):

    def __init__(self, citation_in_text, citation, regex, reference):
        self.citation_in_text = citation_in_text
        self.citation = citation
        self.reference = reference
        self.regex = regex
        self.found = self.reference != None

    def __repr__(self):
        return (
            "{found}\n"
            "Citation:         {citation}\n"
            "Citation In Text: {citation_in_text}\n"
            "Reference:        {reference}\n"
        ).format(found="FOUND" if self.found else "NOT FOUND",
                 citation=self.citation,
                 citation_in_text=self.citation_in_text,
                 reference=self.reference)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('-d', dest='docx_path')
    args = parser.parse_args()
    docx_path = args.docx_path

    result = list()
    references = get_references_as_lines(docx_path)
    text = get_text_as_string(docx_path)

    citations = re.findall(REGEX, text)

    for citation in citations:
        try:
            cites = citation.split(';')
            for index, cite in enumerate(cites, 1):
                cite = cite.replace('(', '').replace(')', '').strip()
                if ' et al' in cite:
                    name = cite.split(' et al., ')[0].strip()
                    year = cite.split(' et al., ')[1].strip()
                    regex = '{}.*\({}\).*'.format(name, year)
                    reference = find_reference_for_cite_regex(regex)
                    result.append(CitationResult(citation_in_text=citation,
                                                 citation=cite,
                                                 regex=regex,
                                                 reference=reference))
                else:
                    names_and_year = cite.split(',')
                    if len(names_and_year) == 2:
                        # Bartram & Casimir, 2007
                        year = names_and_year[1].strip()
                        names = names_and_year[0].split('&')
                        name_regex = ''.join(['{}.*'.format(n.strip()) for n in names])
                        regex = name_regex + '\({}\)'.format(year)
                        reference = find_reference_for_cite_regex(regex)
                        result.append(CitationResult(citation_in_text=citation,
                                                     citation=cite,
                                                     regex=regex,
                                                     reference=reference))
                    else:
                        # Lowe, Kroeck, & Sivasubramaniam, 1996
                        names_and_year[-2] = names_and_year[-2].replace('&', '')
                        names = names_and_year[:-1]
                        year = names_and_year[-1].strip()
                        name_regex = ''.join(['{}.*'.format(n.strip()) for n in names])
                        regex = name_regex + '\({}\)'.format(year)
                        reference = find_reference_for_cite_regex(regex)
                        result.append(CitationResult(citation_in_text=citation,
                                                     citation=cite,
                                                     regex=regex,
                                                     reference=reference))
        except Exception as e:
            result.append(CitationResult(citation_in_text=citation,
                                         citation=None,
                                         regex=None,
                                         reference=None))

    for r in result:
        if not r.found:
            print(r)
